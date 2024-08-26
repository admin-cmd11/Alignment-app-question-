import tkinter as tk
from docx import Document
from docx.shared import Pt

#document setup
doc = Document()
font_size = Pt(12)
#window setup
root = tk.Tk()
root.geometry("720x500")
root.title("Question Alignment App For Teachers")
root.resizable(False,False)
#input q
inp_q = tk.Label(root,text ="Enter \nthe\n question:")
inp_q.place(x = 20,y = 30)
inp_q.config(font=("Berlin Sans FB", 20))
q = tk.Text(root,height=6,width = 43,font=("Arial Unicode MS",15))
q.place(x=190,y =20)
#input a
inp_a = tk.Label(root,text ="a)",font=("Berlin Sans FB",20))
inp_a.place(x=20,y=200)
a = tk.Text(root,height=2,width = 30,font=("Arial Unicode MS",15))
a.place(x=50,y =200)
#input b
inp_b = tk.Label(root,text ="b)",font=("Berlin Sans FB",20))
inp_b.place(x=20,y=260)
b= tk.Text(root,height=2,width = 30,font=("Arial Unicode MS",15))
b.place(x=50,y =260)
#input c
inp_c = tk.Label(root,text ="c)",font=("Berlin Sans FB",20))
inp_c.place(x=20,y=320)
c= tk.Text(root,height=2,width = 30,font=("Arial Unicode MS",15))
c.place(x=50,y =320)
#input d
inp_b = tk.Label(root,text ="d)",font=("Berlin Sans FB",20))
inp_b.place(x=20,y=380)
d= tk.Text(root,height=2,width = 30,font=("Arial Unicode MS",15))
d.place(x=50,y =380)
#storing and covertion to word
def get_entry():
    question = q.get("1.0","end-1c")
    option_a = a.get("1.0","end-1c")
    option_b = b.get("1.0","end-1c")
    option_c = c.get("1.0","end-1c")
    option_d = d.get("1.0","end-1c")
    question_w = doc.add_paragraph()
    run1 = question_w.add_run(question)
    run1.font.size = font_size
    option_a_and_c = doc.add_paragraph()
    aandcoption=("                 "+"a)"+option_a,"                 "+"c)"+option_c)
    run2 = option_a_and_c.add_run(aandcoption)
    run2.font.size = font_size
    option_b_and_d = doc.add_paragraph()
    banddoption =("                 "+"b)"+option_b+"                 "+"d)"+option_d)
    run3 = option_b_and_d.add_run(banddoption)
    run3.font.size = font_size
    doc.save('output.docx')
    
#button
button=tk.Button(text = "DONE",command = get_entry,font = ("Berlin Sans FB",20))
button.place(x =500,y = 290)

#window mainloop
root.mainloop()


